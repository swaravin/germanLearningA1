#!/usr/bin/env python3
"""
German A1 learn pack: Word doc (pronunciation + sentences) + 3 MP3 audio courses.
- German_A1_Learn.docx
- German_A1_Audio_German_and_English.mp3   (DE word → pause → EN meaning → pause)
- German_A1_Audio_English_and_German.mp3   (EN meaning → pause → DE word → pause)
- German_A1_Audio_German_only.mp3          (DE word → pause)
"""

from __future__ import annotations

import asyncio
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

import edge_tts
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.path.insert(0, str(Path(__file__).resolve().parent))

from a1.articles import article_for_german, default_example_sentences, german_with_article  # noqa: E402
from build_german_vocab_doc import SECTIONS, english_number, german_number  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets" / "full"
DOC_PATH = ASSETS / "German_A1_Learn.docx"
AUDIO_DE_EN = ASSETS / "German_A1_Audio_German_and_English.mp3"
AUDIO_EN_DE = ASSETS / "German_A1_Audio_English_and_German.mp3"
AUDIO_DE_ONLY = ASSETS / "German_A1_Audio_German_only.mp3"
README_PATH = ASSETS / "German_A1_Audio_README.txt"
TEMP_DIR = ROOT / ".german_vocab_audio_tmp"

VOICE_DE = "de-DE-KatjaNeural"
VOICE_EN = "en-US-JennyNeural"
RATE_DE = "-32%"  # slow, clear German word
RATE_EN = "+0%"  # natural speed for English meaning (not “pronunciation practice”)
PAUSE_AFTER_WORD_MS = 1.0  # seconds between vocabulary items
PAUSE_DE_BEFORE_EN_MS = 0.9  # gap after German word, before English meaning

try:
    import imageio_ffmpeg

    FFMPEG = imageio_ffmpeg.get_ffmpeg_exe()
except Exception:
    FFMPEG = "ffmpeg"


@dataclass
class VocabEntry:
    section: str
    index: int
    german: str
    english: str
    pronunciation: str
    sentence_de: str
    sentence_en: str


def xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def reading_guide(word: str) -> str:
    """Syllable-style reading guide (hyphenate long words)."""
    w = word.strip()
    if not w or " " in w:
        return w
    if len(w) <= 4:
        return w
    # Simple syllable chunks for learners
    parts = re.findall(r"[A-Za-zÄÖÜäöüß]+|[^A-Za-zÄÖÜäöüß]+", w)
    chunks: list[str] = []
    for p in parts:
        if re.fullmatch(r"[A-Za-zÄÖÜäöüß]+", p) and len(p) > 5:
            mid = len(p) // 2
            chunks.append(p[:mid] + "·" + p[mid:])
        else:
            chunks.append(p)
    return "".join(chunks)


def english_short(en: str) -> str:
    return en.split(" / ")[0].split(" (")[0].strip()


def make_sentences(section: str, de: str, en: str) -> tuple[str, str]:
    return default_example_sentences(de, en, section=section)


def collect_entries() -> list[VocabEntry]:
    entries: list[VocabEntry] = []
    idx = 1
    for section_title, pairs in SECTIONS:
        if section_title.startswith("6."):
            for n in range(101):
                de = german_number(n)
                en = english_number(n)
                sde, sen = make_sentences(section_title, de, en)
                entries.append(
                    VocabEntry(
                        section=section_title,
                        index=idx,
                        german=de,
                        english=en,
                        pronunciation=reading_guide(de),
                        sentence_de=sde,
                        sentence_en=sen,
                    )
                )
                idx += 1
            continue
        for de, en in pairs:
            sde, sen = make_sentences(section_title, de, en)
            entries.append(
                VocabEntry(
                    section=section_title,
                    index=idx,
                    german=de,
                    english=en,
                    pronunciation=reading_guide(de),
                    sentence_de=sde,
                    sentence_en=sen,
                )
            )
            idx += 1
    return entries


def build_doc(entries: list[VocabEntry]) -> None:
    doc = Document()
    title = doc.add_heading("German A1 — Learn Pack", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Each word: reading guide, example sentence in German and English. "
        "Use the two MP3 files for slow, clear pronunciation."
    )

    current_section = ""
    table = None
    first_section = True
    for e in entries:
        if e.section != current_section:
            current_section = e.section
            if not first_section:
                doc.add_page_break()
            first_section = False
            doc.add_heading(e.section, level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            labels = [
                "#",
                "German",
                "Pronunciation (read slowly)",
                "English",
                "Example (German)",
                "Example (English)",
            ]
            for i, lab in enumerate(labels):
                hdr[i].text = lab
                for run in hdr[i].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)

        assert table is not None
        row = table.add_row().cells
        row[0].text = str(e.index)
        row[1].text = e.german
        row[2].text = e.pronunciation
        row[3].text = e.english
        row[4].text = e.sentence_de
        row[5].text = e.sentence_en

    doc.save(DOC_PATH)


def ssml_de_only(e: VocabEntry) -> str:
    de = german_with_article(e.german, section=e.section)
    return (
        "<speak>"
        f'<voice name="{VOICE_DE}">'
        f'<prosody rate="{RATE_DE}">{xml_escape(de)}</prosody>'
        '<break time="1100ms"/>'
        f'<prosody rate="{RATE_DE}">{xml_escape(e.sentence_de)}</prosody>'
        "</voice>"
        '<break time="1800ms"/>'
        "</speak>"
    )


def ssml_de_en(e: VocabEntry) -> str:
    en_word = english_short(e.english)
    de = german_with_article(e.german, section=e.section)
    return (
        "<speak>"
        f'<voice name="{VOICE_DE}">'
        f'<prosody rate="{RATE_DE}">{xml_escape(de)}</prosody>'
        '<break time="1100ms"/>'
        f'<prosody rate="{RATE_DE}">{xml_escape(e.sentence_de)}</prosody>'
        "</voice>"
        '<break time="1600ms"/>'
        f'<voice name="{VOICE_EN}">'
        f'<prosody rate="{RATE_EN}">English: {xml_escape(en_word)}</prosody>'
        '<break time="900ms"/>'
        f'<prosody rate="{RATE_EN}">{xml_escape(e.sentence_en)}</prosody>'
        "</voice>"
        '<break time="2000ms"/>'
        "</speak>"
    )


async def tts_save(text: str, voice: str, rate: str, path: Path) -> None:
    communicate = edge_tts.Communicate(text, voice, rate=rate)
    await communicate.save(str(path))


async def render_entry_audio(
    e: VocabEntry, sem: asyncio.Semaphore
) -> tuple[Path, Path, Path]:
    """
    German-only:     [slow DE word] → pause
    German+English:  [slow DE word] → pause → [EN meaning] → pause
    English+German:  [EN meaning] → pause → [slow DE word] → pause
    """
    async with sem:
        TEMP_DIR.mkdir(parents=True, exist_ok=True)
        path_de = TEMP_DIR / f"{e.index:04d}_de.mp3"
        path_de_en = TEMP_DIR / f"{e.index:04d}_de_en.mp3"
        path_en_de = TEMP_DIR / f"{e.index:04d}_en_de.mp3"
        en_meaning = english_short(e.english)
        de_spoken = german_with_article(e.german, section=e.section)

        path_de_word = TEMP_DIR / f"{e.index:04d}_de_word.mp3"
        path_en_meaning = TEMP_DIR / f"{e.index:04d}_en_meaning.mp3"
        await tts_save(de_spoken, VOICE_DE, RATE_DE, path_de_word)
        await tts_save(en_meaning, VOICE_EN, RATE_EN, path_en_meaning)

        pause_after = ensure_silence_mp3(
            PAUSE_AFTER_WORD_MS, f"pause_after_{PAUSE_AFTER_WORD_MS}s.mp3"
        )
        pause_mid = ensure_silence_mp3(
            PAUSE_DE_BEFORE_EN_MS, f"pause_mid_{PAUSE_DE_BEFORE_EN_MS}s.mp3"
        )

        _concat_files([path_de_word, pause_after], path_de)
        _concat_files(
            [path_de_word, pause_mid, path_en_meaning, pause_after],
            path_de_en,
        )
        _concat_files(
            [path_en_meaning, pause_mid, path_de_word, pause_after],
            path_en_de,
        )

        path_de_word.unlink(missing_ok=True)
        path_en_meaning.unlink(missing_ok=True)
        return path_de, path_de_en, path_en_de


def _concat_files(parts: list[Path], out: Path) -> None:
    list_file = TEMP_DIR / f"concat_{out.name}.txt"
    list_file.write_text(
        "\n".join(f"file '{p.resolve()}'" for p in parts),
        encoding="utf-8",
    )
    subprocess.run(
        [
            FFMPEG,
            "-y",
            "-f",
            "concat",
            "-safe",
            "0",
            "-i",
            str(list_file),
            "-c",
            "copy",
            str(out),
        ],
        check=True,
        capture_output=True,
    )


def concat_mp3(paths: list[Path], out: Path) -> None:
    existing = [p for p in paths if p.exists()]
    if not existing:
        raise RuntimeError("No audio segments to merge.")
    list_file = TEMP_DIR / f"concat_{out.stem}.txt"
    lines = ["file '" + str(p.resolve()).replace("'", "'\\''") + "'" for p in existing]
    list_file.write_text("\n".join(lines), encoding="utf-8")
    subprocess.run(
        [
            FFMPEG,
            "-y",
            "-f",
            "concat",
            "-safe",
            "0",
            "-i",
            str(list_file),
            "-c",
            "copy",
            str(out),
        ],
        check=True,
        capture_output=True,
    )
    print(f"Exported {out} ({len(existing)} segments)")


def ensure_silence_mp3(seconds: float, name: str) -> Path:
    path = TEMP_DIR / name
    if path.exists():
        return path
    subprocess.run(
        [
            FFMPEG,
            "-y",
            "-f",
            "lavfi",
            "-i",
            "anullsrc=r=24000:cl=mono",
            "-t",
            str(seconds),
            "-q:a",
            "9",
            "-acodec",
            "libmp3lame",
            str(path),
        ],
        check=True,
        capture_output=True,
    )
    return path


async def build_audio(entries: list[VocabEntry]) -> None:
    import shutil

    if TEMP_DIR.exists():
        shutil.rmtree(TEMP_DIR)
    TEMP_DIR.mkdir(parents=True)
    sem = asyncio.Semaphore(4)
    de_paths: list[Path] = []
    de_en_paths: list[Path] = []
    en_de_paths: list[Path] = []

    print(f"Generating TTS for {len(entries)} words…")
    batch = 25
    for start in range(0, len(entries), batch):
        chunk = entries[start : start + batch]
        tasks = [render_entry_audio(e, sem) for e in chunk]
        results = await asyncio.gather(*tasks)
        for p_de, p_de_en, p_en_de in results:
            de_paths.append(p_de)
            de_en_paths.append(p_de_en)
            en_de_paths.append(p_en_de)
        print(f"  TTS done {min(start + batch, len(entries))}/{len(entries)}")

    print("Building German-only MP3…")
    concat_mp3(de_paths, AUDIO_DE_ONLY)
    print("Building German → English MP3…")
    concat_mp3(de_en_paths, AUDIO_DE_EN)
    print("Building English → German MP3…")
    concat_mp3(en_de_paths, AUDIO_EN_DE)


def write_readme(entries: list[VocabEntry]) -> None:
    README_PATH.write_text(
        f"""German A1 Audio — How to use (including Spotify)

Files in this folder:
  • {DOC_PATH.name} — vocabulary with sentences
  • {AUDIO_DE_ONLY.name} — slow German word, then pause
  • {AUDIO_DE_EN.name} — German word → pause → English meaning → pause
  • {AUDIO_EN_DE.name} — English meaning → pause → German word → pause

Total vocabulary items: {len(entries)}
Approximate length: listen to the MP3 properties in your player.

─── Spotify ───
Spotify cannot upload arbitrary MP3s as music tracks without a distributor.
You CAN play these files on Spotify Desktop as LOCAL FILES:

1. Install Spotify desktop app (Mac/Windows).
2. Settings → Local Files → turn ON "Show Local Files".
3. Add folder: {ROOT}
4. All three MP3 files appear under "Local Files" in Your Library.
5. Add them to a playlist for learning on phone (download playlist with local sync).

Alternative: Upload as a podcast via Spotify for Creators (podcasters.spotify.com)
— one episode per MP3 or split by section.

─── Phone / car ───
Copy the MP3s to your phone (Files, Google Drive, Apple Music "add file").

─── Learning order ───
German→English: hear the German word, then the English meaning.
English→German: hear the English meaning first, then the German word (good for recall).
Use German_only to test yourself without English hints.
""",
        encoding="utf-8",
    )


async def main_async() -> None:
    import sys

    audio_only = "--audio-only" in sys.argv
    entries = collect_entries()
    print(f"Collected {len(entries)} entries.")
    if not audio_only:
        print("Writing Word document…")
        build_doc(entries)
        write_readme(entries)
    print("Building audio (DE word → pause → EN meaning → pause)…")
    await build_audio(entries)
    if not audio_only:
        write_readme(entries)
    print("Done.")


def main() -> None:
    asyncio.run(main_async())


if __name__ == "__main__":
    main()
