#!/usr/bin/env python3
"""Build German A1 vocabulary Word document with English equivalents and optional images."""

from __future__ import annotations

import io
import time
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parent.parent / "assets" / "full" / "German_A1_Vocabulary.docx"

# German -> English Wikipedia title for image lookup (concrete nouns only)
IMAGE_LOOKUP: dict[str, str] = {
    "Mutter": "Mother",
    "Vater": "Father",
    "Kind": "Child",
    "Haus": "House",
    "Wohnung": "Apartment",
    "Küche": "Kitchen",
    "Tisch": "Table",
    "Stuhl": "Chair",
    "Bett": "Bed",
    "Auto": "Car",
    "Bus": "Bus",
    "Zug": "Train",
    "Fahrrad": "Bicycle",
    "Flugzeug": "Aircraft",
    "Brot": "Bread",
    "Käse": "Cheese",
    "Milch": "Milk",
    "Wasser": "Water",
    "Kaffee": "Coffee",
    "Apfel": "Apple",
    "Banane": "Banana",
    "Orange": "Orange_(fruit)",
    "Tomate": "Tomato",
    "Hund": "Dog",
    "Katze": "Cat",
    "Buch": "Book",
    "Handy": "Mobile_phone",
    "Computer": "Computer",
    "Schule": "School",
    "Krankenhaus": "Hospital",
    "Restaurant": "Restaurant",
    "Hotel": "Hotel",
    "Park": "Park",
    "Kirche": "Church",
    "Supermarkt": "Supermarket",
    "Bahnhof": "Train_station",
    "Flughafen": "Airport",
    "Garten": "Garden",
    "Motorrad": "Motorcycle",
    "Butter": "Butter",
    "Fisch": "Fish",
    "Kino": "Cinema",
    "Museum": "Museum",
    "Bibliothek": "Library",
    "Apotheke": "Pharmacy",
    "Café": "Cafe",
    "Koffer": "Suitcase",
}

SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "1. Personal Pronouns",
        [
            ("ich", "I"),
            ("du", "you (informal singular)"),
            ("er", "he"),
            ("sie", "she / they"),
            ("es", "it"),
            ("wir", "we"),
            ("ihr", "you (informal plural)"),
            ("Sie", "you (formal)"),
        ],
    ),
    (
        "2. Articles & Possessives",
        [
            ("der", "the (masculine)"),
            ("die", "the (feminine/plural)"),
            ("das", "the (neuter)"),
            ("ein", "a/an (masculine/neuter)"),
            ("eine", "a/an (feminine)"),
            ("kein", "no/not a (masculine/neuter)"),
            ("keine", "no/not a (feminine/plural)"),
            ("mein", "my (masculine/neuter)"),
            ("dein", "your (informal, masc./neut.)"),
            ("sein", "his/its (masculine/neuter)"),
            ("ihr", "her/their/your (formal poss.)"),
            ("unser", "our"),
            ("euer", "your (informal plural)"),
        ],
    ),
    (
        "3. Basic Everyday Words",
        [
            ("ja", "yes"),
            ("nein", "no"),
            ("bitte", "please / you're welcome"),
            ("danke", "thank you"),
            ("Entschuldigung", "excuse me / sorry"),
            ("hallo", "hello"),
            ("tschüss", "bye"),
            ("guten Morgen", "good morning"),
            ("guten Tag", "good day / hello"),
            ("gute Nacht", "good night"),
            ("okay", "okay"),
            ("vielleicht", "maybe"),
            ("gut", "good"),
            ("schlecht", "bad"),
            ("sehr", "very"),
            ("viel", "much / a lot"),
            ("wenig", "little / few"),
        ],
    ),
    (
        "4. Question Words",
        [
            ("wer", "who"),
            ("was", "what"),
            ("wo", "where"),
            ("wann", "when"),
            ("warum", "why"),
            ("wie", "how"),
            ("wie viel", "how much"),
            ("wie viele", "how many"),
            ("welcher", "which (masculine)"),
            ("welche", "which (feminine/plural)"),
            ("welches", "which (neuter)"),
        ],
    ),
    (
        "5. Prepositions",
        [
            ("in", "in"),
            ("an", "at/on (vertical surface)"),
            ("auf", "on"),
            ("unter", "under"),
            ("über", "over / about"),
            ("neben", "next to"),
            ("zwischen", "between"),
            ("vor", "in front of / before"),
            ("hinter", "behind"),
            ("mit", "with"),
            ("ohne", "without"),
            ("für", "for"),
            ("von", "from / of"),
            ("zu", "to"),
            ("bei", "at/near/with"),
            ("nach", "after / to (direction)"),
            ("seit", "since / for (time)"),
            ("gegen", "against"),
            ("durch", "through"),
        ],
    ),
    ("6. Numbers (0–100)", []),
    (
        "7. Essential Verbs",
        [
            ("sein", "to be"),
            ("haben", "to have"),
            ("werden", "to become"),
            ("können", "can / to be able to"),
            ("müssen", "must / to have to"),
            ("dürfen", "may / to be allowed to"),
            ("sollen", "should"),
            ("wollen", "to want"),
            ("mögen", "to like"),
            ("gehen", "to go / walk"),
            ("kommen", "to come"),
            ("fahren", "to drive / travel"),
            ("fliegen", "to fly"),
            ("laufen", "to run / walk"),
            ("bleiben", "to stay"),
            ("sitzen", "to sit"),
            ("stehen", "to stand"),
            ("sprechen", "to speak"),
            ("reden", "to talk"),
            ("sagen", "to say"),
            ("fragen", "to ask"),
            ("antworten", "to answer"),
            ("erklären", "to explain"),
            ("zeigen", "to show"),
            ("arbeiten", "to work"),
            ("lernen", "to learn"),
            ("studieren", "to study"),
            ("wohnen", "to live (reside)"),
            ("leben", "to live"),
            ("schlafen", "to sleep"),
            ("aufstehen", "to get up"),
            ("essen", "to eat"),
            ("trinken", "to drink"),
            ("kochen", "to cook"),
            ("kaufen", "to buy"),
            ("bezahlen", "to pay"),
            ("verkaufen", "to sell"),
            ("nehmen", "to take"),
            ("geben", "to give"),
            ("bringen", "to bring"),
            ("bekommen", "to receive / get"),
            ("finden", "to find"),
            ("suchen", "to search"),
            ("sehen", "to see"),
            ("hören", "to hear"),
            ("lesen", "to read"),
            ("schreiben", "to write"),
            ("denken", "to think"),
            ("wissen", "to know (fact)"),
            ("kennen", "to know (person/place)"),
            ("öffnen", "to open"),
            ("schließen", "to close"),
            ("warten", "to wait"),
            ("treffen", "to meet"),
            ("besuchen", "to visit"),
            ("spielen", "to play"),
            ("helfen", "to help"),
            ("machen", "to do / make"),
            ("putzen", "to clean"),
            ("waschen", "to wash"),
            ("rufen", "to call / shout"),
            ("reisen", "to travel"),
        ],
    ),
    (
        "8. Family & People",
        [
            ("Mutter", "mother"),
            ("Vater", "father"),
            ("Eltern", "parents"),
            ("Kind", "child"),
            ("Kinder", "children"),
            ("Sohn", "son"),
            ("Tochter", "daughter"),
            ("Bruder", "brother"),
            ("Schwester", "sister"),
            ("Großmutter", "grandmother"),
            ("Großvater", "grandfather"),
            ("Oma", "grandma"),
            ("Opa", "grandpa"),
            ("Familie", "family"),
            ("Mann", "man / husband"),
            ("Frau", "woman / wife"),
            ("Freund", "friend (male) / boyfriend"),
            ("Freundin", "friend (female) / girlfriend"),
            ("Nachbar", "neighbor (male)"),
            ("Kollege", "colleague (male)"),
            ("Kollegin", "colleague (female)"),
            ("Chef", "boss"),
            ("Lehrer", "teacher (male)"),
            ("Lehrerin", "teacher (female)"),
            ("Mensch", "human / person"),
            ("Leute", "people"),
            ("Name", "name"),
            ("Alter", "age"),
            ("Geburtsdatum", "date of birth"),
            ("Adresse", "address"),
        ],
    ),
    (
        "9. Home & Objects",
        [
            ("Haus", "house"),
            ("Wohnung", "apartment"),
            ("Zimmer", "room"),
            ("Küche", "kitchen"),
            ("Bad", "bathroom"),
            ("Toilette", "toilet"),
            ("Wohnzimmer", "living room"),
            ("Schlafzimmer", "bedroom"),
            ("Balkon", "balcony"),
            ("Garten", "garden"),
            ("Garage", "garage"),
            ("Tisch", "table"),
            ("Stuhl", "chair"),
            ("Bett", "bed"),
            ("Sofa", "sofa"),
            ("Schrank", "cupboard / wardrobe"),
            ("Regal", "shelf"),
            ("Tür", "door"),
            ("Fenster", "window"),
            ("Wand", "wall"),
            ("Boden", "floor"),
            ("Dach", "roof"),
            ("Schlüssel", "key"),
            ("Lampe", "lamp"),
            ("Uhr", "clock / watch"),
            ("Bild", "picture"),
            ("Teppich", "carpet / rug"),
            ("Handy", "mobile phone"),
            ("Telefon", "telephone"),
            ("Computer", "computer"),
            ("Laptop", "laptop"),
            ("Tablet", "tablet"),
            ("Buch", "book"),
            ("Zeitung", "newspaper"),
            ("Stift", "pen"),
            ("Papier", "paper"),
            ("Tasche", "bag"),
            ("Geld", "money"),
            ("Karte", "card / map"),
        ],
    ),
    (
        "10. Food & Drinks",
        [
            ("Essen", "food / meal / to eat"),
            ("Frühstück", "breakfast"),
            ("Mittagessen", "lunch"),
            ("Abendessen", "dinner"),
            ("Brot", "bread"),
            ("Butter", "butter"),
            ("Käse", "cheese"),
            ("Milch", "milk"),
            ("Joghurt", "yogurt"),
            ("Wasser", "water"),
            ("Saft", "juice"),
            ("Kaffee", "coffee"),
            ("Tee", "tea"),
            ("Fleisch", "meat"),
            ("Fisch", "fish"),
            ("Reis", "rice"),
            ("Nudeln", "pasta / noodles"),
            ("Kartoffel", "potato"),
            ("Obst", "fruit"),
            ("Gemüse", "vegetables"),
            ("Apfel", "apple"),
            ("Banane", "banana"),
            ("Orange", "orange"),
            ("Tomate", "tomato"),
            ("Salat", "salad / lettuce"),
            ("Zwiebel", "onion"),
            ("Zucker", "sugar"),
            ("Salz", "salt"),
            ("Öl", "oil"),
            ("Suppe", "soup"),
        ],
    ),
    (
        "11. Places",
        [
            ("Stadt", "city"),
            ("Dorf", "village"),
            ("Land", "country / land"),
            ("Straße", "street"),
            ("Supermarkt", "supermarket"),
            ("Markt", "market"),
            ("Laden", "shop"),
            ("Geschäft", "store / business"),
            ("Bank", "bank"),
            ("Post", "post office"),
            ("Schule", "school"),
            ("Universität", "university"),
            ("Krankenhaus", "hospital"),
            ("Arzt", "doctor"),
            ("Apotheke", "pharmacy"),
            ("Hotel", "hotel"),
            ("Restaurant", "restaurant"),
            ("Café", "café"),
            ("Bahnhof", "train station"),
            ("Flughafen", "airport"),
            ("Haltestelle", "stop (bus/tram)"),
            ("Park", "park"),
            ("Kino", "cinema"),
            ("Museum", "museum"),
            ("Bibliothek", "library"),
            ("Kirche", "church"),
        ],
    ),
    (
        "12. Transport & Travel",
        [
            ("Auto", "car"),
            ("Bus", "bus"),
            ("Zug", "train"),
            ("Fahrrad", "bicycle"),
            ("Motorrad", "motorcycle"),
            ("Flugzeug", "airplane"),
            ("Ticket", "ticket"),
            ("Fahrkarte", "fare ticket"),
            ("Reise", "trip / journey"),
            ("Urlaub", "holiday / vacation"),
            ("Gepäck", "luggage"),
            ("Koffer", "suitcase"),
            ("Weg", "way / path"),
            ("Richtung", "direction"),
            ("Stadtplan", "city map"),
            ("Karte", "map / card"),
        ],
    ),
    (
        "13. Time & Dates",
        [
            ("Tag", "day"),
            ("Woche", "week"),
            ("Monat", "month"),
            ("Jahr", "year"),
            ("heute", "today"),
            ("morgen", "tomorrow"),
            ("gestern", "yesterday"),
            ("Morgen", "morning"),
            ("Abend", "evening"),
            ("Nacht", "night"),
            ("Stunde", "hour"),
            ("Minute", "minute"),
            ("Sekunde", "second"),
            ("Montag", "Monday"),
            ("Dienstag", "Tuesday"),
            ("Mittwoch", "Wednesday"),
            ("Donnerstag", "Thursday"),
            ("Freitag", "Friday"),
            ("Samstag", "Saturday"),
            ("Sonntag", "Sunday"),
            ("Januar", "January"),
            ("Februar", "February"),
            ("März", "March"),
            ("April", "April"),
            ("Mai", "May"),
            ("Juni", "June"),
            ("Juli", "July"),
            ("August", "August"),
            ("September", "September"),
            ("Oktober", "October"),
            ("November", "November"),
            ("Dezember", "December"),
        ],
    ),
    (
        "14. Adjectives",
        [
            ("gut", "good"),
            ("schlecht", "bad"),
            ("groß", "big / tall"),
            ("klein", "small"),
            ("alt", "old"),
            ("jung", "young"),
            ("neu", "new"),
            ("teuer", "expensive"),
            ("billig", "cheap"),
            ("schön", "beautiful"),
            ("hässlich", "ugly"),
            ("einfach", "simple / easy"),
            ("schwer", "difficult / heavy"),
            ("wichtig", "important"),
            ("richtig", "correct"),
            ("falsch", "wrong"),
            ("schnell", "fast"),
            ("langsam", "slow"),
            ("warm", "warm"),
            ("kalt", "cold"),
            ("heiß", "hot"),
            ("glücklich", "happy"),
            ("traurig", "sad"),
            ("müde", "tired"),
            ("gesund", "healthy"),
            ("krank", "sick"),
            ("freundlich", "friendly"),
            ("nett", "nice"),
            ("laut", "loud"),
            ("leise", "quiet"),
        ],
    ),
    (
        "15. Useful Adverbs & Words",
        [
            ("immer", "always"),
            ("nie", "never"),
            ("oft", "often"),
            ("selten", "rarely"),
            ("schon", "already"),
            ("noch", "still / yet"),
            ("wieder", "again"),
            ("auch", "also / too"),
            ("nur", "only"),
            ("hier", "here"),
            ("dort", "there"),
            ("zusammen", "together"),
            ("allein", "alone"),
            ("jetzt", "now"),
            ("später", "later"),
            ("dann", "then"),
        ],
    ),
]


def german_number(n: int) -> str:
    ones = [
        "null",
        "eins",
        "zwei",
        "drei",
        "vier",
        "fünf",
        "sechs",
        "sieben",
        "acht",
        "neun",
        "zehn",
        "elf",
        "zwölf",
        "dreizehn",
        "vierzehn",
        "fünfzehn",
        "sechzehn",
        "siebzehn",
        "achtzehn",
        "neunzehn",
    ]
    tens = [
        "",
        "",
        "zwanzig",
        "dreißig",
        "vierzig",
        "fünfzig",
        "sechzig",
        "siebzig",
        "achtzig",
        "neunzig",
    ]
    if n == 100:
        return "hundert"
    if n < 20:
        return ones[n]
    if n % 10 == 0:
        return tens[n // 10]
    unit = ones[n % 10]
    if unit == "eins":
        unit = "ein"
    return f"{unit}und{tens[n // 10]}"


def english_number(n: int) -> str:
    words = [
        "zero",
        "one",
        "two",
        "three",
        "four",
        "five",
        "six",
        "seven",
        "eight",
        "nine",
        "ten",
        "eleven",
        "twelve",
        "thirteen",
        "fourteen",
        "fifteen",
        "sixteen",
        "seventeen",
        "eighteen",
        "nineteen",
    ]
    if n < 20:
        return words[n]
    if n == 100:
        return "one hundred"
    tens = [
        "",
        "",
        "twenty",
        "thirty",
        "forty",
        "fifty",
        "sixty",
        "seventy",
        "eighty",
        "ninety",
    ]
    if n % 10 == 0:
        return tens[n // 10]
    return f"{tens[n // 10]}-{words[n % 10]}"


def fetch_thumbnail(title: str, session: requests.Session) -> bytes | None:
    url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{requests.utils.quote(title)}"
    try:
        r = session.get(url, timeout=8)
        if r.status_code != 200:
            return None
        data = r.json()
        thumb = data.get("thumbnail") or {}
        src = thumb.get("source")
        if not src:
            return None
        img = session.get(src, timeout=10)
        if img.status_code == 200 and img.content:
            return img.content
    except Exception:
        return None
    return None


def add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_heading(title, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_entry_table_header(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "German"
    hdr[1].text = "English"
    hdr[2].text = "Picture"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)
    return table


def add_word_row(
    table,
    german: str,
    english: str,
    session: requests.Session,
    try_image: bool,
) -> None:
    row = table.add_row().cells
    row[0].text = german
    row[1].text = english
    row[2].text = ""
    if try_image and german in IMAGE_LOOKUP:
        data = fetch_thumbnail(IMAGE_LOOKUP[german], session)
        if data:
            try:
                row[2].paragraphs[0].add_run().add_picture(
                    io.BytesIO(data), width=Inches(0.85)
                )
            except Exception:
                row[2].text = "—"
        else:
            row[2].text = "—"
        time.sleep(0.15)
    else:
        row[2].text = ""


def build_document() -> Document:
    doc = Document()
    title = doc.add_heading("German A1 Vocabulary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Essential German words with English equivalents. "
        "Pictures included for selected concrete nouns (from Wikipedia)."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    session = requests.Session()
    session.headers["User-Agent"] = "GermanVocabDocBuilder/1.0 (educational)"

    first = True
    for section_title, pairs in SECTIONS:
        if not first:
            doc.add_page_break()
        first = False
        add_section_heading(doc, section_title)
        table = add_entry_table_header(doc)
        if section_title.startswith("6."):
            for n in range(101):
                add_word_row(
                    table, german_number(n), english_number(n), session, False
                )
            continue
        for de, en in pairs:
            try_image = section_title.split(".", 1)[0].strip() in {
                "8",
                "9",
                "10",
                "11",
                "12",
            }
            add_word_row(table, de, en, session, try_image)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
